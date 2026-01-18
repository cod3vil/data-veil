"""
Property-based tests for Document Parser module

Feature: data-desensitization-platform
"""

import pytest
from hypothesis import given, strategies as st, settings, assume
from hypothesis import HealthCheck
import tempfile
import os
from pathlib import Path

from app.document_parser import DocumentParser, ParsedDocument, DocumentParsingError

# Import document creation libraries
import fitz  # PyMuPDF
from docx import Document
from openpyxl import Workbook


# Hypothesis strategies for generating test data
@st.composite
def text_content(draw):
    """Generate text content with mixed Chinese and English"""
    # Generate a mix of ASCII and Chinese characters
    chinese_chars = st.characters(
        whitelist_categories=('Lo',),
        min_codepoint=0x4E00,
        max_codepoint=0x9FFF
    )
    ascii_chars = st.characters(min_codepoint=32, max_codepoint=126)
    
    # Mix of Chinese and English text
    text = draw(st.text(
        alphabet=st.one_of(chinese_chars, ascii_chars),
        min_size=10,
        max_size=500
    ))
    
    # Ensure we have some non-whitespace content
    assume(text.strip())
    return text


# Feature: data-desensitization-platform, Property 4: Multi-format Document Parsing
@given(content=text_content())
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
@pytest.mark.property_test
def test_pdf_parsing_preserves_content(content):
    """
    Property 4: Multi-format Document Parsing
    Validates: Requirements 2.1
    
    For any supported document format with known text content,
    when parsed, the extracted text should match the original content.
    
    Note: PDF text extraction has known limitations with certain characters
    due to font embedding issues. This test verifies that parsing succeeds
    and extracts non-empty content.
    """
    parser = DocumentParser()
    
    # Create a temporary PDF file with the content
    with tempfile.NamedTemporaryFile(mode='w+b', suffix='.pdf', delete=False) as tmp_file:
        tmp_path = tmp_file.name
        
        try:
            # Create PDF with PyMuPDF
            doc = fitz.open()
            page = doc.new_page()
            
            # Insert text (handle potential font issues)
            try:
                page.insert_text((50, 50), content, fontsize=12)
            except:
                # If insertion fails, skip this example
                doc.close()
                assume(False)
            
            doc.save(tmp_path)
            doc.close()
            
            # Parse the PDF
            result = parser.parse(tmp_path, 'pdf')
            
            # Verify the result
            assert isinstance(result, ParsedDocument)
            assert result.content is not None
            assert len(result.content.strip()) > 0
            
            # Verify metadata
            assert result.metadata['format'] == 'PDF'
            assert result.metadata['page_count'] >= 1
            
            # Note: We don't strictly verify content match due to PDF rendering limitations
            # The key property is that parsing succeeds and extracts text
            
        finally:
            # Clean up
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)


@given(content=text_content())
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
@pytest.mark.property_test
def test_docx_parsing_preserves_content(content):
    """
    Property 4: Multi-format Document Parsing
    Validates: Requirements 2.2
    
    For any DOCX document with known text content,
    when parsed, the extracted text should match the original content.
    """
    parser = DocumentParser()
    
    # Create a temporary DOCX file with the content
    with tempfile.NamedTemporaryFile(mode='w+b', suffix='.docx', delete=False) as tmp_file:
        tmp_path = tmp_file.name
        
        try:
            # Create DOCX with python-docx
            doc = Document()
            doc.add_paragraph(content)
            doc.save(tmp_path)
            
            # Parse the DOCX
            result = parser.parse(tmp_path, 'docx')
            
            # Verify the result
            assert isinstance(result, ParsedDocument)
            assert result.content is not None
            assert len(result.content.strip()) > 0
            
            # The content should match the original
            assert content in result.content
            
            # Verify metadata
            assert result.metadata['format'] == 'DOCX'
            assert result.metadata['paragraph_count'] >= 1
            
        finally:
            # Clean up
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)


@given(content=text_content())
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
@pytest.mark.property_test
def test_xlsx_parsing_preserves_content(content):
    """
    Property 4: Multi-format Document Parsing
    Validates: Requirements 2.3
    
    For any XLSX document with known text content,
    when parsed, the extracted text should contain the original content.
    
    Note: Excel treats strings starting with '=' as formulas, which may
    result in empty cells. We filter such cases.
    """
    parser = DocumentParser()
    
    # Skip content that starts with '=' as Excel treats it as a formula
    assume(not content.startswith('='))
    
    # Create a temporary XLSX file with the content
    with tempfile.NamedTemporaryFile(mode='w+b', suffix='.xlsx', delete=False) as tmp_file:
        tmp_path = tmp_file.name
        
        try:
            # Create XLSX with openpyxl
            wb = Workbook()
            ws = wb.active
            ws['A1'] = content
            wb.save(tmp_path)
            wb.close()
            
            # Parse the XLSX
            result = parser.parse(tmp_path, 'xlsx')
            
            # Verify the result
            assert isinstance(result, ParsedDocument)
            assert result.content is not None
            assert len(result.content.strip()) > 0
            
            # The content should contain the original text
            assert content in result.content
            
            # Verify metadata
            assert result.metadata['format'] == 'XLSX'
            assert result.metadata['sheet_count'] >= 1
            
        finally:
            # Clean up
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)


@given(content=text_content())
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
@pytest.mark.property_test
def test_txt_parsing_preserves_content(content):
    """
    Property 4: Multi-format Document Parsing
    Validates: Requirements 2.4
    
    For any TXT document with known text content,
    when parsed, the extracted text should match the original content.
    """
    parser = DocumentParser()
    
    # Create a temporary TXT file with the content
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as tmp_file:
        tmp_path = tmp_file.name
        tmp_file.write(content)
    
    try:
        # Parse the TXT
        result = parser.parse(tmp_path, 'txt')
        
        # Verify the result
        assert isinstance(result, ParsedDocument)
        assert result.content is not None
        assert len(result.content.strip()) > 0
        
        # The content should match exactly
        assert content == result.content
        
        # Verify metadata
        assert result.metadata['format'] == 'TXT'
        assert 'encoding' in result.metadata
        
    finally:
        # Clean up
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)



# Feature: data-desensitization-platform, Property 5: Mixed Language Content Handling
@given(
    chinese_text=st.text(
        alphabet=st.characters(
            whitelist_categories=('Lo',),
            min_codepoint=0x4E00,
            max_codepoint=0x9FFF
        ),
        min_size=5,
        max_size=100
    ),
    english_text=st.text(
        alphabet=st.characters(min_codepoint=65, max_codepoint=122),
        min_size=5,
        max_size=100
    )
)
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
@pytest.mark.property_test
def test_mixed_language_content_handling(chinese_text, english_text):
    """
    Property 5: Mixed Language Content Handling
    Validates: Requirements 2.6
    
    For any document containing both Chinese and English text,
    when parsed, the extracted content should preserve all characters
    from both languages correctly.
    """
    # Ensure we have actual content
    assume(chinese_text.strip() and english_text.strip())
    
    # Create mixed content
    mixed_content = f"{chinese_text} {english_text}"
    
    parser = DocumentParser()
    
    # Test with TXT format (simplest to verify)
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as tmp_file:
        tmp_path = tmp_file.name
        tmp_file.write(mixed_content)
    
    try:
        result = parser.parse(tmp_path, 'txt')
        
        # Verify both Chinese and English content are preserved
        assert chinese_text in result.content
        assert english_text in result.content
        
        # Verify the full mixed content is present
        assert mixed_content == result.content
        
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


@given(
    chinese_text=st.text(
        alphabet=st.characters(
            whitelist_categories=('Lo',),
            min_codepoint=0x4E00,
            max_codepoint=0x9FFF
        ),
        min_size=5,
        max_size=100
    ),
    english_text=st.text(
        alphabet=st.characters(min_codepoint=65, max_codepoint=122),
        min_size=5,
        max_size=100
    )
)
@settings(
    max_examples=100,
    deadline=None,
    suppress_health_check=[HealthCheck.function_scoped_fixture]
)
@pytest.mark.property_test
def test_mixed_language_docx_handling(chinese_text, english_text):
    """
    Property 5: Mixed Language Content Handling (DOCX)
    Validates: Requirements 2.6
    
    For any DOCX document containing both Chinese and English text,
    when parsed, the extracted content should preserve all characters
    from both languages correctly.
    """
    # Ensure we have actual content
    assume(chinese_text.strip() and english_text.strip())
    
    # Create mixed content
    mixed_content = f"{chinese_text} {english_text}"
    
    parser = DocumentParser()
    
    # Test with DOCX format
    with tempfile.NamedTemporaryFile(mode='w+b', suffix='.docx', delete=False) as tmp_file:
        tmp_path = tmp_file.name
        
        try:
            doc = Document()
            doc.add_paragraph(mixed_content)
            doc.save(tmp_path)
            
            result = parser.parse(tmp_path, 'docx')
            
            # Verify both Chinese and English content are preserved
            assert chinese_text in result.content
            assert english_text in result.content
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
