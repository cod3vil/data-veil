"""
Document Parser Module

This module provides functionality to parse various document formats
and extract text content for desensitization processing.
"""

from typing import Dict, Any, Optional, List
from dataclasses import dataclass
from abc import ABC, abstractmethod
import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook
import chardet

from app.exceptions import DocumentParsingError
from app.logging_config import get_logger

logger = get_logger(__name__)


@dataclass
class ParsedDocument:
    """Data model for parsed document content"""
    content: str
    metadata: Dict[str, Any]
    structure: Optional[Dict] = None


class DocumentParser:
    """
    Document parser for extracting text content from various file formats.
    
    Supports: PDF, DOCX, XLSX, TXT
    """
    
    def parse(self, file_path: str, file_type: str) -> ParsedDocument:
        """
        Parse document and extract text content.
        
        Args:
            file_path: Path to the document file
            file_type: Type of the document (pdf, docx, xlsx, txt)
            
        Returns:
            ParsedDocument containing extracted content and metadata
            
        Raises:
            DocumentParsingError: If parsing fails
        """
        file_type = file_type.lower()
        logger.info("starting_document_parsing", file_path=file_path, file_type=file_type)
        
        try:
            if file_type == 'pdf':
                result = self.parse_pdf(file_path)
            elif file_type == 'docx':
                result = self.parse_docx(file_path)
            elif file_type == 'xlsx':
                result = self.parse_xlsx(file_path)
            elif file_type == 'txt':
                result = self.parse_txt(file_path)
            else:
                raise DocumentParsingError(
                    f"Unsupported file type: {file_type}",
                    error_code="UNSUPPORTED_FORMAT"
                )
            
            logger.info(
                "document_parsing_complete",
                file_type=file_type,
                content_length=len(result.content),
                metadata=result.metadata
            )
            return result
            
        except DocumentParsingError:
            logger.error(
                "document_parsing_failed",
                file_type=file_type,
                error_code=getattr(DocumentParsingError, 'error_code', 'UNKNOWN')
            )
            raise
        except Exception as e:
            logger.error(
                "document_parsing_unexpected_error",
                file_type=file_type,
                error=str(e)
            )
            raise DocumentParsingError(
                f"Failed to parse {file_type} document: {str(e)}",
                error_code="PARSING_FAILED",
                details={"original_error": str(e)}
            )
    
    def parse_pdf(self, file_path: str) -> ParsedDocument:
        """
        Parse PDF document using PyMuPDF.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            ParsedDocument with extracted text content
            
        Raises:
            DocumentParsingError: If PDF parsing fails
        """
        try:
            doc = fitz.open(file_path)
            
            # Check if document is empty
            if doc.page_count == 0:
                raise DocumentParsingError(
                    "PDF document is empty",
                    error_code="EMPTY_DOCUMENT"
                )
            
            # Extract text from all pages
            text_content = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():  # Only add non-empty pages
                    text_content.append(text)
            
            # Combine all text
            full_text = "\n".join(text_content)
            
            # Extract metadata before closing
            page_count = doc.page_count
            metadata = {
                "page_count": page_count,
                "format": "PDF",
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "subject": doc.metadata.get("subject", ""),
            }
            
            doc.close()
            
            if not full_text.strip():
                raise DocumentParsingError(
                    "No text content found in PDF",
                    error_code="NO_CONTENT"
                )
            
            return ParsedDocument(
                content=full_text,
                metadata=metadata,
                structure={"pages": page_count}
            )
            
        except fitz.FileDataError as e:
            raise DocumentParsingError(
                f"Corrupted or invalid PDF file: {str(e)}",
                error_code="CORRUPTED_FILE"
            )
        except fitz.FileNotFoundError as e:
            raise DocumentParsingError(
                f"PDF file not found: {str(e)}",
                error_code="FILE_NOT_FOUND"
            )
    
    def parse_docx(self, file_path: str) -> ParsedDocument:
        """
        Parse DOCX document using python-docx.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            ParsedDocument with extracted text content including paragraphs and tables
            
        Raises:
            DocumentParsingError: If DOCX parsing fails
        """
        try:
            doc = Document(file_path)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            # Combine all text
            full_text = "\n".join(text_content)
            
            if not full_text.strip():
                raise DocumentParsingError(
                    "No text content found in DOCX",
                    error_code="NO_CONTENT"
                )
            
            # Extract metadata
            metadata = {
                "format": "DOCX",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            
            # Add core properties if available
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                })
            
            return ParsedDocument(
                content=full_text,
                metadata=metadata,
                structure={
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables)
                }
            )
            
        except Exception as e:
            # Check for common DOCX errors
            if "not a zip file" in str(e).lower() or "bad zipfile" in str(e).lower():
                raise DocumentParsingError(
                    "Corrupted or invalid DOCX file",
                    error_code="CORRUPTED_FILE"
                )
            raise DocumentParsingError(
                f"Failed to parse DOCX: {str(e)}",
                error_code="PARSING_FAILED",
                details={"original_error": str(e)}
            )
    
    def parse_xlsx(self, file_path: str) -> ParsedDocument:
        """
        Parse XLSX document using openpyxl.
        
        Args:
            file_path: Path to the XLSX file
            
        Returns:
            ParsedDocument with extracted cell content from all sheets
            
        Raises:
            DocumentParsingError: If XLSX parsing fails
        """
        try:
            workbook = load_workbook(file_path, data_only=True)
            
            if not workbook.sheetnames:
                raise DocumentParsingError(
                    "XLSX file contains no sheets",
                    error_code="EMPTY_DOCUMENT"
                )
            
            text_content = []
            sheet_info = []
            total_non_empty_cells = 0
            
            # Extract text from all sheets
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                
                # Add sheet name as header
                sheet_text.append(f"=== Sheet: {sheet_name} ===")
                
                # Extract cell values
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell) for cell in row if cell is not None and str(cell).strip()]
                    if row_values:
                        sheet_text.append(" | ".join(row_values))
                        total_non_empty_cells += len(row_values)
                
                if len(sheet_text) > 1:  # More than just the header
                    text_content.extend(sheet_text)
                    sheet_info.append({
                        "name": sheet_name,
                        "rows": sheet.max_row,
                        "columns": sheet.max_column
                    })
            
            # Combine all text
            full_text = "\n".join(text_content)
            
            if not full_text.strip() or total_non_empty_cells == 0:
                raise DocumentParsingError(
                    "No content found in XLSX file",
                    error_code="NO_CONTENT"
                )
            
            # Extract metadata
            metadata = {
                "format": "XLSX",
                "sheet_count": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
            }
            
            workbook.close()
            
            return ParsedDocument(
                content=full_text,
                metadata=metadata,
                structure={"sheets": sheet_info}
            )
            
        except Exception as e:
            # Check for common XLSX errors
            if "not a zip file" in str(e).lower() or "bad zipfile" in str(e).lower():
                raise DocumentParsingError(
                    "Corrupted or invalid XLSX file",
                    error_code="CORRUPTED_FILE"
                )
            raise DocumentParsingError(
                f"Failed to parse XLSX: {str(e)}",
                error_code="PARSING_FAILED",
                details={"original_error": str(e)}
            )
    
    def parse_txt(self, file_path: str) -> ParsedDocument:
        """
        Parse TXT document with encoding detection.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            ParsedDocument with extracted text content
            
        Raises:
            DocumentParsingError: If TXT parsing fails
        """
        try:
            # First, detect the encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            
            if not raw_data:
                raise DocumentParsingError(
                    "TXT file is empty",
                    error_code="EMPTY_DOCUMENT"
                )
            
            # Detect encoding
            detection_result = chardet.detect(raw_data)
            detected_encoding = detection_result['encoding']
            confidence = detection_result['confidence']
            
            # Try UTF-8 first (most common for modern files)
            try:
                text_content = raw_data.decode('utf-8')
                encoding = 'utf-8'
            except UnicodeDecodeError:
                # UTF-8 failed, try detected encoding
                if detected_encoding:
                    try:
                        text_content = raw_data.decode(detected_encoding)
                        encoding = detected_encoding
                    except (UnicodeDecodeError, LookupError):
                        # Detected encoding failed, try other common encodings
                        for fallback_encoding in ['gbk', 'gb2312', 'latin-1']:
                            try:
                                text_content = raw_data.decode(fallback_encoding)
                                encoding = fallback_encoding
                                break
                            except (UnicodeDecodeError, LookupError):
                                continue
                        else:
                            raise DocumentParsingError(
                                "Unable to decode TXT file with any known encoding",
                                error_code="ENCODING_ERROR"
                            )
                else:
                    # No detected encoding, try common encodings
                    for fallback_encoding in ['gbk', 'gb2312', 'latin-1']:
                        try:
                            text_content = raw_data.decode(fallback_encoding)
                            encoding = fallback_encoding
                            break
                        except (UnicodeDecodeError, LookupError):
                            continue
                    else:
                        raise DocumentParsingError(
                            "Unable to decode TXT file with any known encoding",
                            error_code="ENCODING_ERROR"
                        )
            
            if not text_content.strip():
                raise DocumentParsingError(
                    "No text content found in TXT file",
                    error_code="NO_CONTENT"
                )
            
            # Extract metadata
            metadata = {
                "format": "TXT",
                "encoding": encoding,
                "encoding_confidence": confidence,
                "size_bytes": len(raw_data),
                "line_count": len(text_content.splitlines()),
            }
            
            return ParsedDocument(
                content=text_content,
                metadata=metadata,
                structure={"lines": len(text_content.splitlines())}
            )
            
        except FileNotFoundError:
            raise DocumentParsingError(
                f"TXT file not found: {file_path}",
                error_code="FILE_NOT_FOUND"
            )
        except PermissionError:
            raise DocumentParsingError(
                f"Permission denied reading TXT file: {file_path}",
                error_code="PERMISSION_DENIED"
            )
