"""
File Exporter Module

This module provides functionality to export desensitized content
to various file formats.
"""

from typing import Dict, Any, Optional
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook


class FileExportError(Exception):
    """Exception raised when file export fails"""
    def __init__(self, message: str, error_code: str = "EXPORT_ERROR", details: Dict = None):
        self.message = message
        self.error_code = error_code
        self.details = details or {}
        super().__init__(self.message)


class FileExporter:
    """
    File exporter for exporting desensitized content to various formats.
    
    Supports: TXT, MD (Markdown), DOCX, XLSX
    """
    
    def _sanitize_xml_string(self, text: str) -> str:
        """
        Sanitize string to be XML-compatible by removing control characters.
        
        Args:
            text: String to sanitize
            
        Returns:
            Sanitized string safe for XML
        """
        # Remove control characters (0x00-0x1F except tab, newline, carriage return)
        # and other problematic characters
        sanitized = ''.join(
            char for char in text
            if ord(char) >= 0x20 or char in '\t\n\r'
        )
        return sanitized
    
    def export(
        self,
        content: str,
        original_format: str,
        output_format: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Export desensitized content to specified format.
        
        Args:
            content: Desensitized text content to export
            original_format: Original document format (pdf, docx, xlsx, txt)
            output_format: Desired output format (txt, md, docx, xlsx)
            metadata: Optional metadata from original document
            
        Returns:
            Bytes of the exported file
            
        Raises:
            FileExportError: If export fails
        """
        output_format = output_format.lower()
        metadata = metadata or {}
        
        try:
            if output_format == 'txt':
                return self._export_txt(content)
            elif output_format == 'md':
                return self._export_md(content, metadata)
            elif output_format == 'docx':
                return self._export_docx(content, metadata)
            elif output_format == 'xlsx':
                return self._export_xlsx(content, metadata)
            else:
                raise FileExportError(
                    f"Unsupported output format: {output_format}",
                    error_code="UNSUPPORTED_FORMAT"
                )
        except FileExportError:
            raise
        except Exception as e:
            raise FileExportError(
                f"Failed to export to {output_format}: {str(e)}",
                error_code="EXPORT_FAILED",
                details={"original_error": str(e)}
            )
    
    def _export_txt(self, content: str) -> bytes:
        """
        Export content as plain text.
        
        Args:
            content: Text content to export
            
        Returns:
            UTF-8 encoded bytes
        """
        return content.encode('utf-8')
    
    def _export_md(self, content: str, metadata: Dict[str, Any]) -> bytes:
        """
        Export content as Markdown format.
        
        Args:
            content: Text content to export
            metadata: Document metadata
            
        Returns:
            UTF-8 encoded Markdown bytes
        """
        md_lines = []
        
        # Add metadata header if available
        if metadata:
            md_lines.append("---")
            if metadata.get("title"):
                md_lines.append(f"title: {metadata['title']}")
            if metadata.get("author"):
                md_lines.append(f"author: {metadata['author']}")
            md_lines.append(f"generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            md_lines.append("desensitized: true")
            md_lines.append("---")
            md_lines.append("")
        
        # Add title if available
        if metadata.get("title"):
            md_lines.append(f"# {metadata['title']}")
            md_lines.append("")
        
        # Process content
        # Split by lines and format appropriately
        lines = content.split('\n')
        
        for line in lines:
            stripped = line.strip()
            
            # Detect sheet headers from XLSX format
            if stripped.startswith("=== Sheet:") and stripped.endswith("==="):
                sheet_name = stripped.replace("=== Sheet:", "").replace("===", "").strip()
                md_lines.append(f"## {sheet_name}")
                md_lines.append("")
            # Detect table rows (contains |)
            elif " | " in stripped:
                # Format as markdown table
                if not md_lines or not md_lines[-1].startswith("|"):
                    # First row - add header separator
                    md_lines.append(f"| {stripped} |")
                    # Count columns
                    col_count = stripped.count(" | ") + 1
                    md_lines.append("| " + " | ".join(["---"] * col_count) + " |")
                else:
                    md_lines.append(f"| {stripped} |")
            else:
                # Regular paragraph
                if stripped:
                    md_lines.append(stripped)
                    md_lines.append("")
        
        # Join all lines
        md_content = "\n".join(md_lines)
        
        return md_content.encode('utf-8')
    
    def _export_docx(self, content: str, metadata: Dict[str, Any]) -> bytes:
        """
        Export content as DOCX format using python-docx.
        
        Args:
            content: Text content to export
            metadata: Document metadata
            
        Returns:
            DOCX file bytes
        """
        doc = Document()
        
        # Set document properties if metadata available
        if metadata:
            core_props = doc.core_properties
            if metadata.get("title"):
                # Sanitize metadata to remove control characters
                core_props.title = self._sanitize_xml_string(metadata["title"])
            if metadata.get("author"):
                core_props.author = self._sanitize_xml_string(metadata["author"])
            if metadata.get("subject"):
                core_props.subject = self._sanitize_xml_string(metadata["subject"])
        
        # Add title if available
        if metadata.get("title"):
            title = doc.add_heading(self._sanitize_xml_string(metadata["title"]), level=1)
        
        # Process content
        lines = content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Detect sheet headers from XLSX format
            if line.startswith("=== Sheet:") and line.endswith("==="):
                sheet_name = line.replace("=== Sheet:", "").replace("===", "").strip()
                doc.add_heading(sheet_name, level=2)
                i += 1
                continue
            
            # Detect table rows (contains |)
            if " | " in line:
                # Collect all consecutive table rows
                table_rows = []
                while i < len(lines) and " | " in lines[i].strip():
                    row_data = [cell.strip() for cell in lines[i].strip().split(" | ")]
                    table_rows.append(row_data)
                    i += 1
                
                if table_rows:
                    # Determine number of columns
                    max_cols = max(len(row) for row in table_rows)
                    
                    # Create table
                    table = doc.add_table(rows=len(table_rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # Fill table
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_value in enumerate(row_data):
                            if col_idx < max_cols:
                                table.rows[row_idx].cells[col_idx].text = cell_value
                
                continue
            
            # Regular paragraph
            if line:
                paragraph = doc.add_paragraph(line)
                # Set font
                for run in paragraph.runs:
                    run.font.size = Pt(11)
            else:
                # Empty line - add spacing
                doc.add_paragraph()
            
            i += 1
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _export_xlsx(self, content: str, metadata: Dict[str, Any]) -> bytes:
        """
        Export content as XLSX format using openpyxl.
        
        Args:
            content: Text content to export
            metadata: Document metadata
            
        Returns:
            XLSX file bytes
        """
        workbook = Workbook()
        
        # Remove default sheet
        if "Sheet" in workbook.sheetnames:
            default_sheet = workbook["Sheet"]
            workbook.remove(default_sheet)
        
        # Process content
        lines = content.split('\n')
        
        current_sheet = None
        current_row = 1
        
        for line in lines:
            stripped = line.strip()
            
            # Detect sheet headers
            if stripped.startswith("=== Sheet:") and stripped.endswith("==="):
                sheet_name = stripped.replace("=== Sheet:", "").replace("===", "").strip()
                # Create new sheet
                current_sheet = workbook.create_sheet(title=sheet_name[:31])  # Excel limit
                current_row = 1
                continue
            
            # Skip empty lines
            if not stripped:
                continue
            
            # If no sheet created yet, create default
            if current_sheet is None:
                current_sheet = workbook.create_sheet(title="Sheet1")
            
            # Parse row data (split by |)
            if " | " in stripped:
                row_data = [cell.strip() for cell in stripped.split(" | ")]
            else:
                row_data = [stripped]
            
            # Write to sheet
            for col_idx, cell_value in enumerate(row_data, start=1):
                current_sheet.cell(row=current_row, column=col_idx, value=cell_value)
            
            current_row += 1
        
        # If no sheets were created, create a default one with content
        if not workbook.sheetnames:
            sheet = workbook.create_sheet(title="Sheet1")
            lines = content.split('\n')
            for row_idx, line in enumerate(lines, start=1):
                if line.strip():
                    sheet.cell(row=row_idx, column=1, value=line.strip())
        
        # Save to BytesIO
        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def generate_filename(
        self,
        original_filename: str,
        output_format: str,
        timestamp: Optional[datetime] = None
    ) -> str:
        """
        Generate filename for exported file following naming convention.
        
        Args:
            original_filename: Original file name
            output_format: Output format extension
            timestamp: Optional timestamp (defaults to now)
            
        Returns:
            Generated filename with format: original_desensitized_YYYYMMDD_HHMMSS.ext
        """
        if timestamp is None:
            timestamp = datetime.now()
        
        # Remove extension from original filename
        if '.' in original_filename:
            base_name = original_filename.rsplit('.', 1)[0]
        else:
            base_name = original_filename
        
        # Format timestamp
        time_str = timestamp.strftime('%Y%m%d_%H%M%S')
        
        # Generate filename
        filename = f"{base_name}_desensitized_{time_str}.{output_format}"
        
        return filename
